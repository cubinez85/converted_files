import os
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
import argparse
from pathlib import Path
import tempfile

# --- Новые импорты для PDF ---
import pdfplumber
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch


# Проверка и установки зависимостей
def check_dependencies():
    try:
        import pandas as pd
        import openpyxl
        import xlrd
        from docx import Document
        import pdfplumber
        import fitz
        import reportlab
    except ImportError as e:
        print(f"Ошибка: отсутствует библиотека: {e.name}")
        print("Установите необходимые зависимости командой:")
        print("pip install pandas openpyxl xlrd python-docx pdfplumber PyMuPDF reportlab")
        exit(1)


class DocumentContent:
    """
    Класс для хранения извлеченного содержимого: текста и таблиц.
    """

    def __init__(self):
        self.text = ""
        self.tables = []  # Список DataFrame'ов

    def add_text(self, text):
        if self.text:
            self.text += "\n" + text
        else:
            self.text = text

    def add_table(self, df):
        self.tables.append(df)

    def has_tables(self):
        return len(self.tables) > 0

    def has_text(self):
        return bool(self.text.strip())

    def to_dict(self):
        """
        Преобразует содержимое в словарь для JSON.
        """
        result = {}
        if self.has_text():
            result['text'] = self.text
        if self.has_tables():
            result['tables'] = [df.to_dict(orient='records') for df in self.tables]
        return result


def extract_content_from_docx(file_path):
    """
    Извлекает текст и таблицы из .docx файла.
    Возвращает объект DocumentContent.
    """
    content = DocumentContent()
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    content.add_text('\n'.join(full_text))

    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            df = pd.DataFrame(data[1:], columns=data[0])  # Первая строка - заголовки
            content.add_table(df)
    return content


def extract_content_from_pdf(file_path):
    """
    Извлекает текст и таблицы из .pdf файла.
    Возвращает объект DocumentContent.
    """
    content = DocumentContent()
    full_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Извлечение текста
            text = page.extract_text()
            if text:
                full_text.append(text)
            # Извлечение таблиц
            tab = page.extract_tables()
            for table in tab:
                if table:
                    df = pd.DataFrame(table[1:], columns=table[0])  # Первая строка - заголовки
                    content.add_table(df)
    content.add_text('\n'.join(full_text))
    return content


def load_data(source_path, target_format):
    """
    Загружает данные из файла в зависимости от расширения.
    Возвращает объект DocumentContent, pandas.DataFrame, список DataFrame, или путь к файлу (для docx -> pdf).
    """
    ext = Path(source_path).suffix.lower()

    if ext == '.csv':
        df = pd.read_csv(source_path)
        content = DocumentContent()
        content.add_table(df)
        return content
    elif ext in ['.xlsx', '.xls']:
        # Возвращает словарь {sheet_name: DataFrame}
        dfs = pd.read_excel(source_path, sheet_name=None)
        content = DocumentContent()
        for name, df in dfs.items():
            content.add_table(df)
        return content
    elif ext == '.json':
        # Предполагаем, что JSON содержит табличные данные
        df = pd.read_json(source_path)
        content = DocumentContent()
        content.add_table(df)
        return content
    elif ext == '.docx':
        if target_format == 'pdf':
            # Для конвертации в PDF возвращаем путь к файлу
            return source_path
        return extract_content_from_docx(source_path)
    elif ext == '.pdf':
        if target_format in ['word', 'excel']:
            return extract_content_from_pdf(source_path)
        else:
            raise ValueError("PDF может быть конвертирован только в Word или Excel.")
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")


def docx_to_pdf_reportlab(docx_path, pdf_path):
    """
    Конвертирует .docx в .pdf с использованием reportlab.
    Извлекает текст и таблицы.
    """
    doc = Document(docx_path)
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    styles = getSampleStyleSheet()
    styleN = styles['Normal']
    styleH = styles['Heading1']

    y_position = height - 40  # Начальная позиция сверху

    for element in doc.element.body:
        if element.tag.endswith('p'):  # Параграф
            text = ''.join([t.text for t in element.xpath('.//w:t')])
            if text.strip():
                p = Paragraph(text, styleN)
                w, h = p.wrap(width - 80, y_position - 20)
                if h > y_position:
                    c.showPage()
                    y_position = height - 40
                p.drawOn(c, 40, y_position - h)
                y_position -= h + 10
        elif element.tag.endswith('tbl'):  # Таблица
            # Извлекаем таблицу через нашу функцию
            temp_content = extract_content_from_docx(docx_path)
            for df in temp_content.tables:
                data = [df.columns.tolist()] + df.values.tolist()
                col_widths = [(width - 80) / len(data[0])] * len(data[0])  # Упрощенная ширина

                table = Table(data, colWidths=col_widths)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                table_width, table_height = table.wrapOn(c, width, height)
                if table_height > y_position:
                    c.showPage()
                    y_position = height - 40
                table.drawOn(c, 40, y_position - table_height)
                y_position -= table_height + 20

    c.save()
    print(f"Конвертировано в PDF: {pdf_path}")


def save_data(data, target_format, output_path):
    """
    Сохраняет данные в нужный формат.
    """
    dir_path = Path(output_path).parent
    os.makedirs(dir_path, exist_ok=True)

    if isinstance(data, DocumentContent):
        # Обработка объекта DocumentContent
        if target_format == 'csv':
            if data.has_tables():
                for i, df in enumerate(data.tables):
                    p = output_path.with_name(f"{output_path.stem}_table_{i + 1}.csv")
                    df.to_csv(p, index=False)
                print(f"Сохранено {len(data.tables)} CSV файлов в {dir_path}")
            else:
                print("Файл не содержит таблиц, невозможно сохранить в CSV.")
        elif target_format == 'json':
            json_data = data.to_dict()
            output_file = output_path.with_suffix('.json')
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            print(f"Сохранено в {output_file}")
        elif target_format == 'excel':
            if data.has_tables():
                output_file = output_path.with_suffix('.xlsx')
                with pd.ExcelWriter(output_file) as writer:
                    for i, df in enumerate(data.tables):
                        df.to_excel(writer, sheet_name=f'Table_{i + 1}', index=False)
                print(f"Сохранено в {output_file}")
            else:
                print("Файл не содержит таблиц, невозможно сохранить в Excel.")
        elif target_format == 'word':
            output_file = output_path.with_suffix('.docx')
            doc = Document()
            if data.has_text():
                doc.add_paragraph(data.text)
                if data.has_tables():
                    doc.add_paragraph('')  # Пустая строка перед таблицей
            if data.has_tables():
                for i, df in enumerate(data.tables):
                    doc.add_heading(f'Таблица {i + 1}', level=1)
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, col_name in enumerate(df.columns):
                        hdr_cells[j].text = str(col_name)
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, value in enumerate(row):
                            row_cells[j].text = str(value)
            doc.save(output_file)
            print(f"Сохранено в {output_file}")
        elif target_format == 'pdf':
            # Это не должно сюда попасть, так как docx -> pdf обрабатывается отдельно
            print("Неверный путь выполнения для Word -> PDF.")
    elif isinstance(data, str) and data.endswith('.docx') and target_format == 'pdf':
        # Специальная ветка для Word -> PDF
        output_pdf = output_path.with_suffix('.pdf')
        docx_to_pdf_reportlab(data, output_pdf)
    elif target_format == 'pdf':
        # Обработка данных (DocumentContent) в PDF
        if isinstance(data, DocumentContent):
            output_file = output_path.with_suffix('.pdf')
            c = canvas.Canvas(str(output_file), pagesize=A4)
            width, height = A4
            styles = getSampleStyleSheet()
            styleN = styles['Normal']
            styleH = styles['Heading1']
            y_position = height - 40

            # Печать текста
            if data.has_text():
                p = Paragraph(data.text, styleN)
                w, h = p.wrap(width - 80, y_position - 20)
                while h > y_position:
                    c.showPage()
                    y_position = height - 40
                    w, h = p.wrap(width - 80, y_position - 20)
                p.drawOn(c, 40, y_position - h)
                y_position -= h + 10

            # Печать таблиц
            if data.has_tables():
                if data.has_text():
                    # Добавляем отступ, если был текст
                    y_position -= 20
                for i, df in enumerate(data.tables):
                    if y_position < 100:  # Проверяем место на странице
                        c.showPage()
                        y_position = height - 40
                    c.setFont("Helvetica-Bold", 12)
                    title = f"Таблица {i + 1}"
                    c.drawString(40, y_position, title)
                    y_position -= 20

                    table_data = [df.columns.tolist()] + df.values.tolist()
                    col_width = (width - 80) / len(table_data[0]) if table_data[0] else 1
                    col_widths = [col_width] * len(table_data[0])

                    table = Table(table_data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))

                    table_width, table_height = table.wrapOn(c, width, height)
                    if table_height > y_position:
                        c.showPage()
                        y_position = height - 40
                    table.drawOn(c, 40, y_position - table_height)
                    y_position -= table_height + 20

            c.save()
            print(f"Сохранено в {output_file}")
        else:
            print(f"Невозможно конвертировать тип {type(data)} в PDF.")
    else:
        print(f"Неподдерживаемая комбинация данных и формата: {type(data)} -> {target_format}")


def main():
    parser = argparse.ArgumentParser(description='Конвертер форматов данных (Excel, CSV, JSON, Word, PDF)')
    parser.add_argument('input_file', type=str, help='Путь к исходному файлу')
    parser.add_argument('output_format', type=str, choices=['csv', 'json', 'excel', 'word', 'pdf'],
                        help='Целевой формат')
    parser.add_argument('output_folder', type=str, help='Папка для сохранения результата')

    args = parser.parse_args()

    check_dependencies()

    input_path = Path(args.input_file)
    output_folder = Path(args.output_folder)
    target_format = args.output_format

    if not input_path.exists():
        print(f"Ошибка: файл не найден - {input_path}")
        return

    try:
        print(f"Загрузка данных из {input_path}...")
        data = load_data(input_path, target_format)

        output_name = input_path.stem
        output_path = output_folder / output_name

        print(f"Конвертация в формат {target_format}...")
        save_data(data, target_format, output_path)

    except ValueError as e:
        print(f"Ошибка: {e}")
    except Exception as e:
        print(f"Произошла ошибка: {e}")


if __name__ == '__main__':
    main()